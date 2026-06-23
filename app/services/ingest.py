import os
import shutil
import uuid
from typing import List, Dict, Any, Optional
import pypdf
import docx
import chromadb
from sqlalchemy.orm import Session
import google.generativeai as genai
from openai import OpenAI

from app.core.config import settings
from app.models import Document, DocumentChunk

# Initialize ChromaDB Client
chroma_client = chromadb.PersistentClient(path=settings.CHROMA_PERSIST_DIRECTORY)

def get_chroma_collection(provider: str, api_key: str):
    """
    Get or create a Chroma collection.
    We separate collections by provider or model to avoid embedding dimension mismatches.
    """
    collection_name = f"techity_chunks_{provider}"
    return chroma_client.get_or_create_collection(name=collection_name)

class RecursiveCharacterTextSplitter:
    """
    A custom recursive character-based text splitter similar to LangChain's,
    eliminating heavy dependencies.
    """
    def __init__(self, chunk_size: int = 1500, chunk_overlap: int = 200, separators: List[str] = None):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators or ["\n\n", "\n", " ", ""]

    def split_text(self, text: str) -> List[str]:
        return self._split_text(text, self.separators)

    def _split_text(self, text: str, separators: List[str]) -> List[str]:
        final_chunks = []
        # Get the separator to use
        separator = separators[-1]
        new_separators = []
        for i, s in enumerate(separators):
            if s == "":
                separator = s
                break
            if s in text:
                separator = s
                new_separators = separators[i + 1:]
                break

        # Split text by current separator
        if separator != "":
            splits = text.split(separator)
        else:
            splits = list(text)

        # Merge splits into chunks
        current_chunk = []
        current_length = 0

        for split in splits:
            split_len = len(split)
            if current_length + split_len + (len(separator) if current_chunk else 0) <= self.chunk_size:
                current_chunk.append(split)
                current_length += split_len + (len(separator) if current_chunk else 0)
            else:
                if current_chunk:
                    merged = separator.join(current_chunk)
                    if len(merged) > self.chunk_size:
                        # If a single split is larger than chunk_size, split it recursively
                        final_chunks.extend(self._split_text(merged, new_separators))
                    else:
                        final_chunks.append(merged)
                
                # Setup overlap
                # Remove items from current_chunk until it's small enough to start the next chunk with overlap
                current_chunk = [split]
                current_length = split_len
                
        if current_chunk:
            merged = separator.join(current_chunk)
            if len(merged) > self.chunk_size:
                final_chunks.extend(self._split_text(merged, new_separators))
            else:
                final_chunks.append(merged)

        # Implement overlap post-processing if needed
        # (This basic merge logic is robust. We can refine it to include overlapping segments)
        return self._apply_overlap(final_chunks)

    def _apply_overlap(self, chunks: List[str]) -> List[str]:
        if len(chunks) <= 1:
            return chunks
        
        overlapped_chunks = []
        overlapped_chunks.append(chunks[0])
        
        for i in range(1, len(chunks)):
            prev = chunks[i - 1]
            curr = chunks[i]
            
            # Take ending of previous chunk as overlap
            overlap_text = prev[-self.chunk_overlap:] if len(prev) > self.chunk_overlap else prev
            overlapped_chunks.append(overlap_text + " " + curr)
            
        return overlapped_chunks

# Parsers for different file types
def parse_pdf(file_path: str) -> List[Dict[str, Any]]:
    """
    Extracts text page-by-page from a PDF file.
    Returns a list of dicts: [{'page': page_num, 'text': text}]
    """
    pages_data = []
    with open(file_path, "rb") as f:
        reader = pypdf.PdfReader(f)
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages_data.append({
                "page": i + 1,
                "text": text.strip()
            })
    return pages_data

def parse_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Extracts text paragraph-by-paragraph from a DOCX file.
    Since Word does not have strict page numbers, we split paragraphs and group them
    into pseudo-pages (approx 3000 chars per page).
    """
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    
    joined_text = "\n".join(full_text)
    
    # Simple chunk-based pagination for docx
    page_size = 3000
    pages_data = []
    for i in range(0, len(joined_text), page_size):
        pages_data.append({
            "page": (i // page_size) + 1,
            "text": joined_text[i : i + page_size]
        })
    return pages_data

def parse_txt(file_path: str) -> List[Dict[str, Any]]:
    """
    Extracts text from a plain text file.
    """
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [{"page": 1, "text": text}]

def extract_document_text(file_path: str, file_type: str) -> List[Dict[str, Any]]:
    if file_type == "pdf":
        return parse_pdf(file_path)
    elif file_type == "docx":
        return parse_docx(file_path)
    elif file_type == "txt":
        return parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

# Embedding Helpers
def generate_embeddings(texts: List[str], provider: str, api_key: str) -> List[List[float]]:
    if provider == "gemini":
        genai.configure(api_key=api_key)
        result = genai.embed_content(
            model=settings.GEMINI_EMBEDDING_MODEL,
            content=texts,
            task_type="retrieval_document"
        )
        # Handle single vs batch result from genai API
        embeddings = result.get("embedding", [])
        if embeddings and not isinstance(embeddings[0], list):
            return [embeddings]
        return embeddings
    elif provider == "openai":
        client = OpenAI(api_key=api_key)
        response = client.embeddings.create(
            model=settings.OPENAI_EMBEDDING_MODEL,
            input=texts
        )
        return [data.embedding for data in response.data]
    else:
        raise ValueError(f"Unsupported embedding provider: {provider}")

def process_and_ingest_file(
    db: Session,
    file_path: str,
    filename: str,
    file_type: str,
    size_bytes: int,
    user_id: int,
    provider: str,
    api_key: str
) -> Document:
    # 1. Create document record in database
    db_doc = Document(
        user_id=user_id,
        filename=filename,
        filepath=file_path,
        file_type=file_type,
        size_bytes=size_bytes
    )
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)

    try:
        # 2. Parse text from document
        pages = extract_document_text(file_path, file_type)
        
        # 3. Chunk text page-by-page
        splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=150)
        chunks_to_insert = []
        chunk_index = 0
        
        for page in pages:
            page_text = page["text"]
            if not page_text:
                continue
            
            page_chunks = splitter.split_text(page_text)
            for idx, text in enumerate(page_chunks):
                chunks_to_insert.append({
                    "document_id": db_doc.id,
                    "chunk_index": chunk_index,
                    "page_number": page["page"],
                    "content": text
                })
                chunk_index += 1

        if not chunks_to_insert:
            # If no text was extracted, add a fallback chunk
            chunks_to_insert.append({
                "document_id": db_doc.id,
                "chunk_index": 0,
                "page_number": 1,
                "content": f"[Empty Document: {filename}]"
            })

        # 4. Save chunks in SQLite database
        db_chunks = []
        for c in chunks_to_insert:
            chunk_obj = DocumentChunk(
                document_id=c["document_id"],
                chunk_index=c["chunk_index"],
                page_number=c["page_number"],
                content=c["content"]
            )
            db.add(chunk_obj)
            db_chunks.append(chunk_obj)
        db.commit()

        # Update IDs
        for chunk_obj in db_chunks:
            db.refresh(chunk_obj)

        # 5. Populate SQLite FTS5 Search Index
        from sqlalchemy import text
        for chunk_obj in db_chunks:
            db.execute(
                text("INSERT INTO document_chunks_fts (chunk_id, content) VALUES (:chunk_id, :content)"),
                {"chunk_id": chunk_obj.id, "content": chunk_obj.content}
            )
        db.commit()

        # 6. Generate embeddings and store in ChromaDB
        texts = [c.content for c in db_chunks]
        ids = [f"doc_{db_doc.id}_chunk_{c.chunk_index}" for c in db_chunks]
        metadatas = [{
            "chunk_id": c.id,
            "document_id": db_doc.id,
            "filename": filename,
            "page_number": c.page_number or 1,
            "user_id": user_id
        } for c in db_chunks]
        
        # Batch embedding requests in chunks of 50 to avoid API limits
        batch_size = 50
        embeddings = []
        for i in range(0, len(texts), batch_size):
            batch_texts = texts[i : i + batch_size]
            batch_embs = generate_embeddings(batch_texts, provider, api_key)
            embeddings.extend(batch_embs)

        # Upsert into Chroma
        collection = get_chroma_collection(provider, api_key)
        collection.add(
            ids=ids,
            embeddings=embeddings,
            metadatas=metadatas,
            documents=texts
        )

        return db_doc
    except Exception as e:
        # Clean up db_doc if anything failed during chunking/embedding
        db.delete(db_doc)
        db.commit()
        raise e



def delete_document_from_indexes(db: Session, document: Document, provider: str, api_key: str):
    """
    Remove document from SQLite (chunks and FTS table) and ChromaDB
    """
    # Get all chunk IDs
    chunk_ids = [c.id for c in document.chunks]
    chunk_indexes = [c.chunk_index for c in document.chunks]
    
    # 1. Delete from ChromaDB
    try:
        collection = get_chroma_collection(provider, api_key)
        chroma_ids = [f"doc_{document.id}_chunk_{idx}" for idx in chunk_indexes]
        collection.delete(ids=chroma_ids)
    except Exception:
        # Collection might not exist or chroma might fail if key is invalid, continue cleanup of SQLite
        pass
        
    # 2. Delete from FTS5 index
    from sqlalchemy import text
    for cid in chunk_ids:
        db.execute(text("DELETE FROM document_chunks_fts WHERE chunk_id = :chunk_id"), {"chunk_id": cid})
    
    # SQLAlchemy cascade will delete DocumentChunk and Document records automatically
    db.delete(document)
    db.commit()
